from PIL import Image
import pytesseract
import whisper
from pypdf import PdfReader
from docx import Document
import os
import platform
import re
import unicodedata
from typing import List, Tuple

def configure_tesseract():
    """Tự động cấu hình đường dẫn Tesseract OCR"""
    if hasattr(pytesseract.pytesseract, 'tesseract_cmd') and pytesseract.pytesseract.tesseract_cmd:
        configured_path = pytesseract.pytesseract.tesseract_cmd
        if configured_path and os.path.exists(configured_path):
            return  
    
    from dotenv import load_dotenv
    load_dotenv()
    env_path = os.getenv('TESSERACT_CMD')
    
    if env_path:
        env_path = os.path.normpath(env_path.strip().strip('"').strip("'"))
        
        paths_to_try = [env_path]
        if platform.system() == 'Windows' and '/' in env_path:
            paths_to_try.append(env_path.replace('/', '\\'))
        
        for test_path in paths_to_try:
            if os.path.exists(test_path):
                pytesseract.pytesseract.tesseract_cmd = test_path
                return
        
        print(f"⚠️  TESSERACT_CMD trong .env không tồn tại: {env_path}")
        print(f"   Đã thử các path: {paths_to_try}")
        print(f"   Hãy kiểm tra lại đường dẫn trong file .env")
    
    if platform.system() == 'Windows':
        username = os.getenv('USERNAME', '')
        possible_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            r'C:\Users\{}\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'.format(username),
            r'D:\Program Files\Tesseract-OCR\tesseract.exe',
            r'D:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        ]
        
        for path in possible_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                return

configure_tesseract()

def process_image_file(file_path: str) -> Tuple[str, str]:
    """
    Extract text từ image bằng OCR
    """
    try:
        configure_tesseract()
        
        if not os.path.exists(file_path):
            error_msg = f"File không tồn tại: {file_path}"
            return '', error_msg
        
        try:
            version = pytesseract.get_tesseract_version()
            print(f"✅ Tesseract OCR version: {version}")
        except Exception as tess_err:
            error_msg = f"Tesseract OCR chưa được cài đặt hoặc không tìm thấy: {tess_err}. Hãy cài Tesseract và set TESSERACT_CMD trong .env"
            return '', error_msg
        
        try:
            image = Image.open(file_path)
        except Exception as img_err:
            error_msg = f"Không thể mở file ảnh: {img_err}"
            return '', error_msg
        try:
            text = pytesseract.image_to_string(image, lang='vie+eng')
            if text.strip():
                return text, ''
            else:
                error_msg = "OCR trả về text rỗng - có thể ảnh không có text, chất lượng ảnh kém, hoặc text quá nhỏ/mờ"
                return '', error_msg
        except Exception as ocr_err:
            print(f"❌ Error khi chạy OCR (vie+eng): {ocr_err}")
            try:
                print("🔄 Thử lại với chỉ tiếng Anh...")
                text = pytesseract.image_to_string(image, lang='eng')
                print(f"✅ OCR (tiếng Anh) thành công, độ dài text: {len(text)} ký tự")
                if text.strip():
                    return text, ''
                else:
                    error_msg = f"OCR (tiếng Anh) trả về text rỗng. Lỗi ban đầu: {ocr_err}"
                    return '', error_msg
            except Exception as ocr_err2:
                error_msg = f"OCR thất bại cả tiếng Việt và tiếng Anh. Lỗi: {ocr_err2}"
                print(f"❌ Error khi chạy OCR (tiếng Anh): {ocr_err2}")
                return '', error_msg
        
    except Exception as e:
        error_msg = f"Lỗi không xác định: {type(e).__name__}: {e}"
        print(f"❌ Error không xác định trong process_image_file: {error_msg}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        return '', error_msg

def process_audio_file(file_path: str) -> str:
    """Transcribe audio thành text bằng Whisper"""
    try:
        model = whisper.load_model('small')
        res = model.transcribe(file_path)
        return res.get('text','')
    except Exception as e:
        return ''

def process_pdf_file(file_path: str) -> str:
    """
    Extract text từ PDF file
    Hỗ trợ PDF có text layer
    """
    try:
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return '\n'.join(text_parts)
    except Exception as e:
        return ''

def process_docx_file(file_path: str) -> str:
    """
    Extract text từ DOCX file
    Hỗ trợ .docx format
    """
    try:
        doc = Document(file_path)
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        
        return '\n'.join(paragraphs)
    except Exception as e:
        return ''

_BULLET_PATTERN = re.compile(r'^\s*[-•●·]\s+', re.MULTILINE)
_MULTI_SPACE_PATTERN = re.compile(r'\s+')
_EXTRA_PUNCT_PATTERN = re.compile(r'([!?.,;:]){2,}')
_SPACE_BEFORE_PUNCT = re.compile(r'\s+([!?.,;:])')
_SENTENCE_BOUNDARY = re.compile(r'(?<=[.!?])\s+')

COMMON_SPELLING_ERRORS = {
    'ko': 'không',
    'k': 'không',
    'kg': 'không',
    'hok': 'không',
    'khg': 'không',
    'đc': 'được',
    'dc': 'được',
    'dc.': 'được',
    'hok.': 'không',
    'mik': 'mình',
    'mk': 'mình',
    'bt': 'bình thường',
    'bh': 'bây giờ',
    'teh': 'the',
    'recieve': 'receive',
    'adress': 'address',
}


def _normalize_unicode(text: str) -> str:
    return unicodedata.normalize('NFC', text)


def _standardize_bullets(text: str) -> str:
    return _BULLET_PATTERN.sub(lambda match: f"\n- ", text)


def _collapse_whitespace(text: str) -> str:
    return _MULTI_SPACE_PATTERN.sub(' ', text)


def _fix_repeated_punctuation(text: str) -> str:
    return _EXTRA_PUNCT_PATTERN.sub(lambda m: m.group(1), text)


def _trim_space_before_punct(text: str) -> str:
    return _SPACE_BEFORE_PUNCT.sub(r'\1', text)


def _basic_spell_correct(text: str) -> str:
    tokens = re.split(r'(\W+)', text)
    corrected: List[str] = []
    for token in tokens:
        key = token.lower()
        if key in COMMON_SPELLING_ERRORS:
            replacement = COMMON_SPELLING_ERRORS[key]
            if token.istitle():
                replacement = replacement.capitalize()
            elif token.isupper():
                replacement = replacement.upper()
            corrected.append(replacement)
        else:
            corrected.append(token)
    return ''.join(corrected)


def _normalize_sentence_spacing(text: str) -> str:
    sentences = [seg.strip() for seg in _SENTENCE_BOUNDARY.split(text) if seg.strip()]
    return ' '.join(sentences)


def clean_text(text: str) -> str:
    """
    Làm sạch, chuẩn hóa text: bỏ khoảng trắng thừa, chuẩn hóa dấu câu,
    sửa các lỗi chính tả phổ biến và nối câu hợp lý.
    """
    if not text:
        return ''
    
    normalized = _normalize_unicode(text)
    normalized = normalized.replace('\r', ' ').strip()
    normalized = _standardize_bullets(normalized)
    normalized = _collapse_whitespace(normalized)
    normalized = _fix_repeated_punctuation(normalized)
    normalized = _trim_space_before_punct(normalized)
    normalized = _basic_spell_correct(normalized)
    normalized = _normalize_sentence_spacing(normalized)
    
    return normalized.strip()
